from datetime import datetime
from openpyxl import Workbook
from docx import Document
from rooms.models import Room, CheckInInformation
from django.db.models import Count, F, ExpressionWrapper, fields
from django.utils import timezone

class ReportGenerator:
    @staticmethod
    def generate_room_occupancy_report(format='XLSX'):
        rooms = Room.objects.annotate(
            occupancy_rate=ExpressionWrapper(
                F('occupied_seats') * 100 / F('seats'),
                output_field=fields.FloatField()
            )
        )

        if format == 'XLSX':
            wb = Workbook()
            ws = wb.active
            ws.title = "Заполняемость комнат"
            
            # Заголовки
            headers = ['Номер комнаты', 'Всего мест', 'Занято мест', 'Процент заполняемости']
            for col, header in enumerate(headers, 1):
                ws.cell(row=1, column=col, value=header)
            
            # Данные
            for row, room in enumerate(rooms, 2):
                ws.cell(row=row, column=1, value=room.id)
                ws.cell(row=row, column=2, value=room.seats)
                ws.cell(row=row, column=3, value=room.occupied_seats)
                ws.cell(row=row, column=4, value=f"{room.occupancy_rate:.2f}%")
            
            return wb

        else:  # DOCX
            doc = Document()
            doc.add_heading('Заполняемость комнат', 0)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Заголовки
            headers = ['Номер комнаты', 'Всего мест', 'Занято мест', 'Процент заполняемости']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            # Данные
            for room in rooms:
                row = table.add_row()
                row.cells[0].text = str(room.id)
                row.cells[1].text = str(room.seats)
                row.cells[2].text = str(room.occupied_seats)
                row.cells[3].text = f"{room.occupancy_rate:.2f}%"
            
            return doc

    @staticmethod
    def generate_student_distribution_report(format='XLSX'):
        students = CheckInInformation.objects.values(
            'university', 'faculty', 'course'
        ).annotate(
            count=Count('id')
        ).order_by('university', 'faculty', 'course')

        if format == 'XLSX':
            wb = Workbook()
            ws = wb.active
            ws.title = "Распределение студентов"
            
            headers = ['Университет', 'Факультет', 'Курс', 'Количество студентов']
            for col, header in enumerate(headers, 1):
                ws.cell(row=1, column=col, value=header)
            
            for row, student in enumerate(students, 2):
                ws.cell(row=row, column=1, value=student['university'])
                ws.cell(row=row, column=2, value=student['faculty'])
                ws.cell(row=row, column=3, value=student['course'])
                ws.cell(row=row, column=4, value=student['count'])
            
            return wb

        else:  # DOCX
            doc = Document()
            doc.add_heading('Распределение студентов', 0)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            headers = ['Университет', 'Факультет', 'Курс', 'Количество студентов']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for student in students:
                row = table.add_row()
                row.cells[0].text = student['university']
                row.cells[1].text = student['faculty']
                row.cells[2].text = str(student['course'])
                row.cells[3].text = str(student['count'])
            
            return doc

    @staticmethod
    def generate_stay_duration_report(format='XLSX'):
        students = CheckInInformation.objects.annotate(
            stay_duration=ExpressionWrapper(
                F('check_out_date') - F('check_in_date'),
                output_field=fields.DurationField()
            )
        )

        if format == 'XLSX':
            wb = Workbook()
            ws = wb.active
            ws.title = "Сроки проживания"
            
            headers = ['ФИО', 'Комната', 'Дата заселения', 'Дата выселения', 'Срок проживания']
            for col, header in enumerate(headers, 1):
                ws.cell(row=1, column=col, value=header)
            
            for row, student in enumerate(students, 2):
                ws.cell(row=row, column=1, value=f"{student.surname} {student.name} {student.fathername}")
                ws.cell(row=row, column=2, value=student.room.id)
                ws.cell(row=row, column=3, value=student.check_in_date.strftime('%Y-%m-%d'))
                ws.cell(row=row, column=4, value=student.check_out_date.strftime('%Y-%m-%d'))
                ws.cell(row=row, column=5, value=str(student.stay_duration.days))
            
            return wb

        else:  # DOCX
            doc = Document()
            doc.add_heading('Сроки проживания', 0)
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            headers = ['ФИО', 'Комната', 'Дата заселения', 'Дата выселения', 'Срок проживания']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
            
            for student in students:
                row = table.add_row()
                row.cells[0].text = f"{student.surname} {student.name} {student.fathername}"
                row.cells[1].text = str(student.room.id)
                row.cells[2].text = student.check_in_date.strftime('%Y-%m-%d')
                row.cells[3].text = student.check_out_date.strftime('%Y-%m-%d')
                row.cells[4].text = str(student.stay_duration.days)
            
            return doc 